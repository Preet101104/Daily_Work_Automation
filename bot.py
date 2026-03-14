import os
import json
from dotenv import load_dotenv
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables from .env file
load_dotenv()

class NewsItem(BaseModel):
    header: str = Field(description="The header/title of the news item")
    summary_lines: List[str] = Field(
        description="Exactly 3 lines/bullet points of summary", 
        min_length=3, max_length=3
    )

class FinancialUpdates(BaseModel):
    executive_overview: List[str] = Field(
        description="4 to 5 bullet points summarizing the overall deal activity, institutional capital, and macro repositioning context based on the news."
    )
    details: List[NewsItem] = Field(
        description="The detailed news items for M&A and ECM."
    )

def create_word_document(updates: FinancialUpdates, output_filename="Financial_Updates.docx"):
    """Creates a styled Word Document similar to the provided template."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Equity Capital Market Updates', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.color.rgb = RGBColor(0x4F, 0x81, 0xBD) # A nice corporate blue
    
    subtitle = doc.add_paragraph('M&A Momentum, AI Bets & Macro Repositioning\n')
    subtitle.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    # Executive Overview
    doc.add_heading('Executive Overview:', level=1)
    for point in updates.executive_overview:
        # docx doesn't perfectly style custom bullets out of the box without template, 
        # so we use generic List Bullet
        doc.add_paragraph(point, style='List Bullet')
        
    doc.add_paragraph('\n') # spacer

    # Details
    doc.add_heading('Details:', level=1)
    
    for item in updates.details:
        # Header (Mimicking the Blue text style)
        h2 = doc.add_heading(item.header, level=2)
        h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        h2.runs[0].font.size = Pt(12)
        
        # Summary lines
        for line in item.summary_lines:
            doc.add_paragraph(line)
            
        doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Disclaimer typically at the end of the document
    doc.add_page_break()
    doc.add_heading('Disclaimer', level=1)
    disclaimer_text = (
        "This research summary is based on publicly available news sources and "
        "is intended solely for informational purposes. While reasonable care has "
        "been taken to ensure accuracy, the information may be incomplete or "
        "subject to change. This document should not be considered as "
        "investment advice or a recommendation to buy or sell any securities."
    )
    doc.add_paragraph(disclaimer_text)

    doc.save(output_filename)
    print(f"Document successfully saved to {output_filename}")

def extract_financial_updates(pdf_path: str):
    """
    Analyzes a newspaper PDF using Gemini 2.5 Flash to extract
    M&A and ECM updates.
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key or api_key == "your_api_key_here":
        print("Error: Please set your GEMINI_API_KEY in the .env file.")
        return

    # Initialize the Gemini GenAI client
    client = genai.Client()

    print(f"Uploading {pdf_path} to Gemini...")
    try:
        # Upload the PDF file directly to Gemini's File API
        uploaded_file = client.files.upload(file=pdf_path)
    except Exception as e:
        print(f"Failed to upload the file: {e}")
        return

    prompt = """
    You are an expert financial analyst. Please carefully read the provided newspaper PDF.
    Identify and extract news and updates strictly related to:
    1. Mergers and Acquisitions (M&A)
    2. Equity Capital Markets (ECM)
    """

    print("Analyzing the document with Gemini 2.5 Flash... This may take a few moments.")
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[uploaded_file, prompt],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=FinancialUpdates,
            ),
        )
        
        # Parse output into pydantic model
        structured_data = FinancialUpdates.model_validate_json(response.text)
        
        # Pass to the python-docx generator
        create_word_document(structured_data, output_filename="Financial_Updates_Output.docx")
        
    except Exception as e:
        print(f"Error generating content: {e}")


if __name__ == "__main__":
    # Ensure you have a sample PDF to test with
    sample_pdf = "newspaper.pdf"
    
    if os.path.exists(sample_pdf):
        extract_financial_updates(sample_pdf)
    else:
        print(f"Please place a PDF named '{sample_pdf}' in the current directory,")
        print("or update the 'sample_pdf' variable in this script with your PDF path.")
